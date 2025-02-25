import os
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def identify_table_type(df):
    """识别财务报表类型"""
    if df.empty or len(df.columns) < 2:
        return None
    
    # 将DataFrame的前几行转换为字符串进行搜索
    header_text = ' '.join([str(x) for x in df.iloc[:5].values.flatten()])
    
    # 定义关键字匹配规则
    patterns = {
        '资产负债表(合并)': ['合并资产负债表', '合并资产负债', '资产负债表'],
        '资产负债表(母公司)': ['母公司资产负债表', '母公司资产负债'],
        '利润表(合并)': ['合并利润表', '合并损益表', '利润表'],
        '利润表(母公司)': ['母公司利润表', '母公司损益表'],
        '现金流量表(合并)': ['合并现金流量表', '合并现金流', '现金流量表'],
        '现金流量表(母公司)': ['母公司现金流量表', '母公司现金流']
    }
    
    # 检查是否包含母公司关键字
    is_parent = '母公司' in header_text
    
    # 根据关键字匹配表格类型
    for table_type, keywords in patterns.items():
        if any(keyword in header_text for keyword in keywords):
            return table_type
        
    # 如果没有明确匹配，尝试通过表格内容特征判断
    if '资产' in header_text and '负债' in header_text:
        return '资产负债表(母公司)' if is_parent else '资产负债表(合并)'
    elif '收入' in header_text and '利润' in header_text:
        return '利润表(母公司)' if is_parent else '利润表(合并)'
    elif '现金' in header_text and '流量' in header_text:
        return '现金流量表(母公司)' if is_parent else '现金流量表(合并)'
    
    return None

def validate_table_data(df, table_type):
    """验证表格数据的完整性"""
    if df is None or df.empty:
        return False
    
    # 定义各类报表必须包含的关键字
    required_keywords = {
        '资产负债表': ['资产总计', '负债总计', '所有者权益'],
        '利润表': ['营业收入', '营业利润', '净利润'],
        '现金流量表': ['经营活动', '投资活动', '筹资活动']
    }
    
    # 获取当前表格类型的基础类型
    base_type = next((k for k in required_keywords.keys() if k in table_type), None)
    if not base_type:
        return False
    
    # 将DataFrame转换为字符串用于检索
    table_text = ' '.join([str(x) for x in df.values.flatten()])
    
    # 检查必需关键字是否存在
    return all(keyword in table_text for keyword in required_keywords[base_type])

def analyze_reports(reports):
    """解析PDF报表，转换为Excel和Word格式"""
    results = []
    errors = []
    
    print(f"开始解析报表，共{len(reports) if reports else 0}个文件")
    
    if not reports or len(reports) == 0:
        raise ValueError("未选择任何报表进行解析")
    
    for report_path in reports:
        try:
            print(f"\n正在处理报表：{report_path}")
            full_path = os.path.join(os.getcwd(), report_path)
            if not os.path.exists(full_path):
                error_msg = f"文件不存在：{report_path}"
                print(f"错误：{error_msg}")
                errors.append(error_msg)
                continue

            if not report_path.lower().endswith('.pdf'):
                error_msg = f"不支持的文件格式：{report_path}，仅支持PDF文件"
                print(f"错误：{error_msg}")
                errors.append(error_msg)
                continue

            print("开始解析PDF文件...")
            doc = fitz.open(full_path)
            
            base_dir = os.path.dirname(full_path)
            file_name = os.path.splitext(os.path.basename(full_path))[0]
            
            result = {
                'title': file_name,
                'excelPath': None,
                'wordPath': None
            }
            
            financial_tables = {
                '资产负债表(合并)': None,
                '资产负债表(母公司)': None,
                '利润表(合并)': None,
                '利润表(母公司)': None,
                '现金流量表(合并)': None,
                '现金流量表(母公司)': None
            }
            text_content = []
            
            for page in doc:
                try:
                    print(f"正在处理第{page.number + 1}页...")
                    # 提取表格，使用更严格的表格识别参数
                    page_tables = page.find_tables(
                        vertical_strategy="lines",
                        horizontal_strategy="lines",
                        snap_tolerance=3,
                        join_tolerance=3,
                        edge_min_length=3
                    )
                    tables_found = list(page_tables)
                    print(f"发现{len(tables_found)}个表格")
                    
                    for table in tables_found:
                        if table.cells and len(table.cells) > 0:
                            table_data = table.extract()
                            if table_data:
                                df = pd.DataFrame(table_data)
                                # 清理数据：移除空行和空列
                                df = df.dropna(how='all').dropna(axis=1, how='all')
                                # 尝试识别表格类型
                                table_type = identify_table_type(df)
                                if table_type and table_type in financial_tables:
                                    # 验证表格数据完整性
                                    if validate_table_data(df, table_type):
                                        financial_tables[table_type] = df
                                    elif financial_tables[table_type] is None:
                                        # 如果当前表格不完整但之前没有数据，也先保存
                                        financial_tables[table_type] = df
                    
                    # 提取文本
                    blocks = page.get_text("blocks")
                    text_blocks = []
                    for block in blocks:
                        is_in_table = False
                        block_rect = fitz.Rect(block[:4])
                        for table in page_tables:
                            table_rect = fitz.Rect(table.bbox)
                            if block_rect.intersects(table_rect):
                                is_in_table = True
                                break
                        if not is_in_table:
                            text_blocks.append(block[4])
                    text_content.extend(text_blocks)
                    print(f"提取了{len(text_blocks)}个文本块")
                    
                except Exception as e:
                    error_msg = f"处理文件 {file_name} 第 {page.number + 1} 页时出错：{str(e)}"
                    print(f"错误：{error_msg}")
                    errors.append(error_msg)
                    continue
            
            # 保存Excel文件
            tables_to_save = {k: v for k, v in financial_tables.items() if v is not None}
            if tables_to_save:
                try:
                    print("正在保存Excel文件...")
                    excel_path = os.path.join(base_dir, f"{file_name}_数据.xlsx")
                    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                        for sheet_name, df in tables_to_save.items():
                            df.to_excel(writer, sheet_name=sheet_name, index=False)
                            worksheet = writer.sheets[sheet_name]
                            # 自动调整列宽
                            for idx, col in enumerate(df.columns):
                                max_length = max(
                                    df[col].astype(str).apply(len).max(),
                                    len(str(col))
                                ) + 2  # 添加一些padding
                                worksheet.column_dimensions[chr(65 + idx)].width = min(max_length, 50)  # 限制最大宽度为50
                    result['excelPath'] = os.path.join(
                        *full_path.split(os.sep)[-4:-1],
                        f"{file_name}_数据.xlsx"
                    )
                    print(f"Excel文件保存成功：{excel_path}")
                except Exception as e:
                    error_msg = f"保存Excel文件时出错 {file_name}：{str(e)}"
                    print(f"错误：{error_msg}")
                    errors.append(error_msg)
            
            # 保存Word文件
            if text_content:
                try:
                    print("正在保存Word文件...")
                    doc_word = Document()
                    title = doc_word.add_heading(file_name, level=1)
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    for text in text_content:
                        if text.strip():
                            p = doc_word.add_paragraph(text.strip())
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            for run in p.runs:
                                run.font.size = Pt(12)
                                run.font.name = '宋体'
                    
                    word_path = os.path.join(base_dir, f"{file_name}_文本.docx")
                    doc_word.save(word_path)
                    result['wordPath'] = os.path.join(
                        *full_path.split(os.sep)[-4:-1],
                        f"{file_name}_文本.docx"
                    )
                    print(f"Word文件保存成功：{word_path}")
                except Exception as e:
                    error_msg = f"保存Word文件时出错 {file_name}：{str(e)}"
                    print(f"错误：{error_msg}")
                    errors.append(error_msg)
            
            results.append(result)
            print(f"报表 {file_name} 处理完成")
            
        except Exception as e:
            error_msg = f"处理文件 {report_path} 时出错：{str(e)}"
            print(f"错误：{error_msg}")
            errors.append(error_msg)
            continue
        finally:
            if 'doc' in locals() and isinstance(doc, fitz.Document):
                doc.close()
    
    if not results and errors:
        error_msg = "\n".join(errors)
        print(f"解析过程中出现错误：\n{error_msg}")
        raise Exception(error_msg)
    
    print(f"\n解析完成，成功处理{len(results)}个文件，失败{len(errors)}个文件")
    return results