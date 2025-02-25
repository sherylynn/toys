import os
import fitz
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from .analyze import identify_table_type, validate_table_data

class BaseProcessor:
    def __init__(self, report_path):
        self.report_path = report_path
        self.full_path = os.path.join(os.getcwd(), report_path)
        self.base_dir = os.path.dirname(self.full_path)
        self.file_name = os.path.splitext(os.path.basename(self.full_path))[0]
        self.doc = None
        self.financial_tables = {
            '资产负债表(合并)': None,
            '资产负债表(母公司)': None,
            '利润表(合并)': None,
            '利润表(母公司)': None,
            '现金流量表(合并)': None,
            '现金流量表(母公司)': None
        }
        self.text_content = []

    def process(self):
        raise NotImplementedError("子类必须实现process方法")

    def save_excel(self):
        tables_to_save = {k: v for k, v in self.financial_tables.items() if v is not None}
        if not tables_to_save:
            return None

        excel_path = os.path.join(self.base_dir, f"{self.file_name}_数据.xlsx")
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            for sheet_name, df in tables_to_save.items():
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                worksheet = writer.sheets[sheet_name]
                for idx, col in enumerate(df.columns):
                    max_length = max(
                        df[col].astype(str).apply(len).max(),
                        len(str(col))
                    ) + 2
                    worksheet.column_dimensions[chr(65 + idx)].width = min(max_length, 50)

        return os.path.join(
            *self.full_path.split(os.sep)[-4:-1],
            f"{self.file_name}_数据.xlsx"
        )

    def save_word(self):
        if not self.text_content:
            return None

        doc_word = Document()
        title = doc_word.add_heading(self.file_name, level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for text in self.text_content:
            if text.strip():
                p = doc_word.add_paragraph(text.strip())
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        word_path = os.path.join(self.base_dir, f"{self.file_name}_文本.docx")
        doc_word.save(word_path)
        return os.path.join(
            *self.full_path.split(os.sep)[-4:-1],
            f"{self.file_name}_文本.docx"
        )

    def cleanup(self):
        if self.doc:
            self.doc.close()

class Doc2xProcessor(BaseProcessor):
    def process(self):
        try:
            self.doc = fitz.open(self.full_path)
            for page in self.doc:
                # 提取表格
                page_tables = page.find_tables(
                    vertical_strategy="lines",
                    horizontal_strategy="lines",
                    snap_tolerance=3,
                    join_tolerance=3,
                    edge_min_length=3
                )
                tables_found = list(page_tables)

                for table in tables_found:
                    if table.cells and len(table.cells) > 0:
                        table_data = table.extract()
                        if table_data:
                            df = pd.DataFrame(table_data)
                            df = df.dropna(how='all').dropna(axis=1, how='all')
                            table_type = identify_table_type(df)
                            if table_type and table_type in self.financial_tables:
                                if validate_table_data(df, table_type) or self.financial_tables[table_type] is None:
                                    self.financial_tables[table_type] = df

                # 提取文本
                blocks = page.get_text("blocks")
                text_blocks = []
                for block in blocks:
                    is_in_table = False
                    block_rect = fitz.Rect(block[:4])
                    for table in page_tables:
                        table_rect = fitz.Rect(table.bbox)
                        if block_rect.intersects(table_rect):
                            is_in_table = True
                            break
                    if not is_in_table:
                        text_blocks.append(block[4])
                self.text_content.extend(text_blocks)

            return {
                'title': self.file_name,
                'excelPath': self.save_excel(),
                'wordPath': self.save_word()
            }

        finally:
            self.cleanup()

class RAGFlowProcessor(BaseProcessor):
    def process(self):
        try:
            self.doc = fitz.open(self.full_path)
            # 使用RAGFlow的增强处理逻辑
            # TODO: 集成RAGFlow的处理逻辑
            # 这里暂时使用与Doc2x相同的处理逻辑
            for page in self.doc:
                page_tables = page.find_tables(
                    vertical_strategy="lines",
                    horizontal_strategy="lines",
                    snap_tolerance=3,
                    join_tolerance=3,
                    edge_min_length=3
                )
                tables_found = list(page_tables)

                for table in tables_found:
                    if table.cells and len(table.cells) > 0:
                        table_data = table.extract()
                        if table_data:
                            df = pd.DataFrame(table_data)
                            df = df.dropna(how='all').dropna(axis=1, how='all')
                            table_type = identify_table_type(df)
                            if table_type and table_type in self.financial_tables:
                                if validate_table_data(df, table_type) or self.financial_tables[table_type] is None:
                                    self.financial_tables[table_type] = df

                blocks = page.get_text("blocks")
                text_blocks = []
                for block in blocks:
                    is_in_table = False
                    block_rect = fitz.Rect(block[:4])
                    for table in page_tables:
                        table_rect = fitz.Rect(table.bbox)
                        if block_rect.intersects(table_rect):
                            is_in_table = True
                            break
                    if not is_in_table:
                        text_blocks.append(block[4])
                self.text_content.extend(text_blocks)

            return {
                'title': self.file_name,
                'excelPath': self.save_excel(),
                'wordPath': self.save_word()
            }

        finally:
            self.cleanup()

def create_processor(report_path, process_mode):
    """根据处理模式创建相应的处理器"""
    if process_mode == 'doc2x':
        return Doc2xProcessor(report_path)
    elif process_mode == 'ragflow':
        return RAGFlowProcessor(report_path)
    else:
        raise ValueError(f"不支持的处理模式：{process_mode}")